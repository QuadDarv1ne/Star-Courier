#!/usr/bin/env python3
"""
Утилита для извлечения текста из .docx файлов
Конвертирует документы в текстовый формат для использования в игре
"""

import os
import sys
from pathlib import Path
from docx import Document


def extract_text_from_docx(filepath: str) -> str:
    """
    Извлечь текст из .docx файла.
    
    Args:
        filepath: Путь к файлу
    
    Returns:
        Извлечённый текст
    """
    try:
        doc = Document(filepath)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Также извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        return "\n\n".join(paragraphs)
    
    except Exception as e:
        return f"[Ошибка чтения файла: {e}]"


def process_directory(input_dir: str, output_dir: str = None):
    """
    Обработать все .docx файлы в директории.
    
    Args:
        input_dir: Входная директория
        output_dir: Выходная директория (по умолчанию создаётся text_output)
    """
    input_path = Path(input_dir)
    
    if not input_path.exists():
        print(f"Директория не найдена: {input_dir}")
        return
    
    if output_dir is None:
        output_dir = str(input_path / "text_output")
    
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    docx_files = list(input_path.glob("*.docx"))
    
    if not docx_files:
        print(f"Нет .docx файлов в директории: {input_dir}")
        return
    
    print(f"Найдено файлов: {len(docx_files)}")
    print()
    
    for docx_file in docx_files:
        print(f"Обработка: {docx_file.name}")
        
        text = extract_text_from_docx(str(docx_file))
        
        # Сохраняем текст в файл с тем же именем, но .txt
        output_file = output_path / f"{docx_file.stem}.txt"
        
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(f"# Файл: {docx_file.name}\n")
            f.write(f"# Извлечено утилитой Star Courier\n")
            f.write("=" * 60 + "\n\n")
            f.write(text)
        
        print(f"  → Сохранено: {output_file.name}")
        print(f"  → Размер: {len(text)} символов")
        print()
    
    print(f"Готово! Тексты сохранены в: {output_dir}")


def process_single_file(filepath: str, output_file: str = None):
    """
    Обработать один файл.
    
    Args:
        filepath: Путь к файлу
        output_file: Путь для выходного файла
    """
    path = Path(filepath)
    
    if not path.exists():
        print(f"Файл не найден: {filepath}")
        return
    
    print(f"Обработка: {path.name}")
    
    text = extract_text_from_docx(str(path))
    
    if output_file is None:
        output_file = str(path.with_suffix(".txt"))
    
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(f"# Файл: {path.name}\n")
        f.write(f"# Извлечено утилитой Star Courier\n")
        f.write("=" * 60 + "\n\n")
        f.write(text)
    
    print(f"  → Сохранено: {output_file}")
    print(f"  → Размер: {len(text)} символов")


def main():
    """Точка входа"""
    print("=" * 60)
    print("  Star Courier — Утилита извлечения текста из .docx")
    print("=" * 60)
    print()
    
    if len(sys.argv) > 1:
        # Режим командной строки
        mode = sys.argv[1]
        
        if mode == "file" and len(sys.argv) > 2:
            filepath = sys.argv[2]
            output = sys.argv[3] if len(sys.argv) > 3 else None
            process_single_file(filepath, output)
        
        elif mode == "dir" and len(sys.argv) > 2:
            dirpath = sys.argv[2]
            output = sys.argv[3] if len(sys.argv) > 3 else None
            process_directory(dirpath, output)
        
        else:
            print("Использование:")
            print("  python extract_docx.py file <путь_к_файлу> [выход_файл]")
            print("  python extract_docx.py dir <путь_к_папке> [выход_папка]")
    else:
        # Интерактивный режим
        print("Выберите режим работы:")
        print("  1. Обработать файл")
        print("  2. Обработать директорию")
        print("  3. Обработать chapters/")
        print("  4. Обработать characters/")
        print()
        
        choice = input("Ваш выбор (1-4): ").strip()
        
        if choice == "1":
            filepath = input("Путь к файлу: ").strip()
            process_single_file(filepath)
        
        elif choice == "2":
            dirpath = input("Путь к директории: ").strip()
            process_directory(dirpath)
        
        elif choice == "3":
            process_directory("chapters")
        
        elif choice == "4":
            process_directory("characters")
        
        else:
            print("Неверный выбор")
    
    print()
    input("Нажмите Enter для выхода...")


if __name__ == "__main__":
    main()
